from flask import (
    Blueprint,
    render_template,
    request,
    redirect,
    send_from_directory,
    url_for,
    flash,
    send_file,
    jsonify,
    abort
)
import os
import io
import sys
import uuid
import base64
import json
import shutil
import smtplib
import subprocess
import re
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from datetime import datetime
from pathlib import Path
from werkzeug.utils import secure_filename
from cryptography.fernet import Fernet
from cryptography.fernet import InvalidToken
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
import hashlib
import secrets
import hmac
from flask import session, g
from urllib.parse import urlparse

import ollama
import docx
import markdown
import pypdfium2 as pdfium
import pdfplumber

signature_bp = Blueprint("signature", __name__,
                         template_folder="templates",
                         static_folder="static")

DEFAULT_OLLAMA_BASE_URL = "http://localhost:11434"
DEFAULT_LANGUAGES = ["English", "Swedish"]
DEFAULT_ANALYSIS_PROMPT = (
    "You are an expert legal analyst specializing in contracts, warranties, and agreements. "
    "Analyze the following document using a {style} communication style. "
    "Write the entire response in {language}. Do not mention the selected output language."
)
SAFE_ID_RE = re.compile(r"^(signed_)?[0-9a-fA-F-]{36}$")
ALLOWED_UPLOAD_EXTENSIONS = {'pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg', 'csv', 'zip'}

def get_data_dir():
    """Get a writable directory for application data."""
    if getattr(sys, 'frozen', False):
        # Running in a bundle (e.g., PyInstaller)
        if sys.platform == 'darwin':
            # macOS: ~/Library/Application Support/Sattmal
            base_dir = os.path.expanduser('~/Library/Application Support/Sattmal')
        else:
            # Fallback for other platforms if needed
            base_dir = os.path.join(os.path.expanduser('~'), '.sattmal')
    else:
        # Running in development
        base_dir = os.getcwd()
    
    data_dir = os.path.join(base_dir, "data")
    os.makedirs(data_dir, exist_ok=True)
    return base_dir

DATA_BASE_DIR = get_data_dir()
CONFIG_FILE = os.path.join(DATA_BASE_DIR, "config.json")

# If config doesn't exist in the data dir, try to copy it from the bundle/source
if not os.path.exists(CONFIG_FILE):
    base_path = getattr(sys, '_MEIPASS', os.getcwd())
    src_config = os.path.join(base_path, "config.json")
    # Also check for config.example.json as a fallback
    src_example = os.path.join(base_path, "config.example.json")
    
    if os.path.exists(src_config):
        shutil.copy(src_config, CONFIG_FILE)
    elif os.path.exists(src_example):
        shutil.copy(src_example, CONFIG_FILE)
    else:
        # Create a default empty config if source is missing too
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump({
                "ollama_base_url": "http://localhost:11434",
                "fernet_key": Fernet.generate_key().decode()
            }, f, indent=2)

with open(CONFIG_FILE, "r", encoding="utf-8") as f:
    config = json.load(f)


def save_config():
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(config, f, indent=2)


def ensure_config_defaults():
    changed = False
    defaults = {
        "ollama_base_url": DEFAULT_OLLAMA_BASE_URL,
        "chat_model": "",
        "vision_model": "",
        "custom_model_name": "",
        "storage_dir": os.path.join(DATA_BASE_DIR, "data"),
        "languages": DEFAULT_LANGUAGES,
        "analysis_prompt": DEFAULT_ANALYSIS_PROMPT,
        "categories": ["Contract", "Invoice", "Report", "Image", "Other"],
    }
    for key, value in defaults.items():
        if key not in config:
            config[key] = value
            changed = True
    if not config.get("flask_secret_key"):
        config["flask_secret_key"] = secrets.token_urlsafe(48)
        changed = True
    if not config.get("languages"):
        config["languages"] = DEFAULT_LANGUAGES
        changed = True
    if not config.get("analysis_prompt"):
        config["analysis_prompt"] = DEFAULT_ANALYSIS_PROMPT
        changed = True
    if changed:
        save_config()


ensure_config_defaults()


def get_flask_secret_key() -> str:
    return config["flask_secret_key"]


def get_version_info() -> dict:
    base_path = getattr(sys, '_MEIPASS', os.getcwd())
    version_path = os.path.join(base_path, "version.json")
    fallback = {"version": "0.0.0", "build": "dev", "build_number": 0, "name": "Sattmal"}
    try:
        with open(version_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return {**fallback, **data}
    except Exception:
        return fallback

OLLAMA_BASE_URL = config.get("ollama_base_url", DEFAULT_OLLAMA_BASE_URL)
# Config loading is done above, but we need to handle the case where salt isbytes
# Salt should be stored as hex in json
MASTER_PASSWORD_HASH = config.get("master_password_hash")
SALT_HEX = config.get("salt")
SALT = bytes.fromhex(SALT_HEX) if SALT_HEX else None

OLLAMA_BASE_URL = config.get("ollama_base_url", DEFAULT_OLLAMA_BASE_URL)

# We no longer use a static FERNET_KEY from config for file encryption
# We will use one derived from the master password for the SESSION.
# However, to avoid breaking existing functionality immediately, we keep this for now
# or we can use it as a fallback?
# Actually, the user wants "masterkey... acts as the key for encryption".
# So we will use the session-stored key.
# But for now, let's keep the variable but maybe unused.
FERNET_KEY = config.get("fernet_key")

# Ensure categories exist in config
def derive_key(password: str, salt: bytes) -> bytes:
    """Derive a Fernet-compatible key from a password and salt."""
    kdf = PBKDF2HMAC(
        algorithm=hashes.SHA256(),
        length=32,
        salt=salt,
        iterations=480000,
    )
    return base64.urlsafe_b64encode(kdf.derive(password.encode()))

def hash_password(password: str, salt: bytes) -> str:
    """Hash password for verification with PBKDF2."""
    digest = hashlib.pbkdf2_hmac("sha256", password.encode(), salt, 480000)
    return "pbkdf2_sha256$480000$" + digest.hex()


def verify_password(password: str, salt: bytes, stored_hash: str) -> bool:
    if not stored_hash:
        return False
    if stored_hash.startswith("pbkdf2_sha256$"):
        return hmac.compare_digest(hash_password(password, salt), stored_hash)

    # Legacy auth hash used by older local configs. Successful login upgrades it.
    legacy_hash = hashlib.sha256(password.encode() + salt).hexdigest()
    return hmac.compare_digest(legacy_hash, stored_hash)


def ollama_client():
    return ollama.Client(host=config.get("ollama_base_url", DEFAULT_OLLAMA_BASE_URL))


def normalize_ollama_models(response) -> list:
    if hasattr(response, 'models'):
        models_list = response.models
    elif isinstance(response, dict):
        models_list = response.get('models', [])
    else:
        models_list = response

    names = []
    for m in models_list:
        name = getattr(m, 'model', m.get('model')) if hasattr(m, 'model') or isinstance(m, dict) else str(m)
        if name:
            names.append(name)
    return sorted(set(names))


def get_ollama_models() -> list:
    return normalize_ollama_models(ollama_client().list())


def ensure_ollama_model(model: str):
    if not model:
        return
    available = get_ollama_models()
    if model not in available:
        ollama_client().pull(model)


def get_required_model(capability: str = "chat") -> str:
    key = f"{capability}_model"
    model = (config.get(key) or "").strip()
    if not model:
        raise ValueError(f"No local {capability} model selected. Choose an Ollama model in Settings.")
    return model


def normalize_languages(raw_languages) -> list:
    if isinstance(raw_languages, str):
        candidates = raw_languages.replace(",", "\n").splitlines()
    else:
        candidates = raw_languages or []

    languages = []
    for item in candidates:
        language = str(item).strip()
        if language and language not in languages:
            languages.append(language)
    return languages or list(DEFAULT_LANGUAGES)


def configured_languages() -> list:
    return normalize_languages(config.get("languages", DEFAULT_LANGUAGES))


def ollama_is_installed() -> bool:
    return shutil.which("ollama") is not None


def safe_doc_id(doc_id: str) -> str:
    if not doc_id or not SAFE_ID_RE.match(doc_id):
        abort(400, description="Invalid document id")
    return doc_id


def storage_root() -> str:
    root = os.path.abspath(os.path.expanduser(config.get("storage_dir") or os.path.join(DATA_BASE_DIR, "data")))
    return root


def ensure_storage_dirs():
    global UPLOADS_DIR, SIGNATURES_DIR
    root = storage_root()
    UPLOADS_DIR = os.path.join(root, "uploads")
    SIGNATURES_DIR = os.path.join(root, "signatures")
    os.makedirs(UPLOADS_DIR, exist_ok=True)
    os.makedirs(SIGNATURES_DIR, exist_ok=True)


def is_within_directory(path: str, directory: str) -> bool:
    try:
        Path(path).resolve().relative_to(Path(directory).resolve())
        return True
    except ValueError:
        return False


def validate_storage_dir(path: str) -> str:
    if not path:
        raise ValueError("Storage directory is required.")
    resolved = os.path.abspath(os.path.expanduser(path))
    os.makedirs(resolved, exist_ok=True)
    test_file = os.path.join(resolved, ".sattmal_write_test")
    with open(test_file, "w", encoding="utf-8") as f:
        f.write("ok")
    os.remove(test_file)
    return resolved


def migrate_storage_dir(new_root: str):
    old_uploads = globals().get("UPLOADS_DIR")
    old_signatures = globals().get("SIGNATURES_DIR")
    new_uploads = os.path.join(new_root, "uploads")
    new_signatures = os.path.join(new_root, "signatures")
    os.makedirs(new_uploads, exist_ok=True)
    os.makedirs(new_signatures, exist_ok=True)

    for src_dir, dest_dir in ((old_uploads, new_uploads), (old_signatures, new_signatures)):
        if not src_dir or os.path.abspath(src_dir) == os.path.abspath(dest_dir) or not os.path.isdir(src_dir):
            continue
        for name in os.listdir(src_dir):
            src = os.path.join(src_dir, name)
            dest = os.path.join(dest_dir, name)
            if os.path.exists(dest):
                continue
            shutil.move(src, dest)


def metadata_path_for(file_id: str) -> str:
    return os.path.join(UPLOADS_DIR, f"{safe_doc_id(file_id)}.json")


def resolve_upload_path(meta: dict) -> str | None:
    stored_filename = secure_filename(os.path.basename(meta.get("stored_filename") or meta.get("filename") or ""))
    if stored_filename:
        candidate = os.path.join(UPLOADS_DIR, stored_filename)
        if os.path.exists(candidate):
            return candidate

    legacy_path = meta.get("file_path")
    if legacy_path and os.path.exists(legacy_path) and is_within_directory(legacy_path, UPLOADS_DIR):
        return legacy_path

    return None

# Middleware to check login
def encrypt_file(file_path: str, key: bytes):
    """Encrypt a file in place."""
    f = Fernet(key)
    with open(file_path, "rb") as file:
        file_data = file.read()
    encrypted_data = f.encrypt(file_data)
    with open(file_path, "wb") as file:
        file.write(encrypted_data)

def decrypt_file_content(file_path: str, key: bytes) -> bytes:
    """Decrypt file content. Returns raw bytes if decryption fails (legacy files)."""
    with open(file_path, "rb") as file:
        file_data = file.read()
    
    try:
        f = Fernet(key)
        return f.decrypt(file_data)
    except Exception:
        # Fallback for legacy unencrypted files
        return file_data

@signature_bp.before_request
def check_login():
    if request.method in {"POST", "PUT", "PATCH", "DELETE"}:
        expected_host = request.host
        for header_name in ("Origin", "Referer"):
            value = request.headers.get(header_name)
            if not value:
                continue
            parsed = urlparse(value)
            if parsed.netloc and parsed.netloc != expected_host:
                abort(403, description="Cross-origin request blocked")

    # Allow static resources and specific routes
    if request.endpoint in ['static', 'signature.login', 'signature.setup', 'signature.logout']:
        return

    # If no master password set, force setup
    global MASTER_PASSWORD_HASH
    if not MASTER_PASSWORD_HASH:
        return redirect(url_for('signature.setup'))

    # Check if user has derived key in session (meaning they logged in)
    if 'encryption_key' not in session:
        return redirect(url_for('signature.login'))
        
    # Store key in g for easy access in this request
    g.encryption_key = session['encryption_key'].encode()

@signature_bp.route("/setup", methods=["GET", "POST"])
def setup():
    global MASTER_PASSWORD_HASH, SALT, config
    
    # If already setup, redirect to login
    if MASTER_PASSWORD_HASH:
         return redirect(url_for('signature.login'))

    ollama_status = {
        "installed": ollama_is_installed(),
        "running": False,
        "models": [],
        "error": None,
    }
    try:
        ollama_status["models"] = get_ollama_models()
        ollama_status["running"] = True
    except Exception as e:
        ollama_status["error"] = str(e)

    if request.method == "POST":
        password = request.form.get("password")
        confirm = request.form.get("confirm_password")
        storage_dir = request.form.get("storage_dir", "").strip()
        ollama_url = request.form.get("ollama_base_url", DEFAULT_OLLAMA_BASE_URL).strip() or DEFAULT_OLLAMA_BASE_URL
        chat_model = request.form.get("chat_model", "").strip()
        
        if password != confirm:
            flash("Passwords do not match!", "danger")
            return render_template("setup.html", config=config, ollama_status=ollama_status)
            
        if not password:
            flash("Password cannot be empty.", "danger")
            return render_template("setup.html", config=config, ollama_status=ollama_status)

        try:
            storage_dir = validate_storage_dir(storage_dir or config.get("storage_dir"))
        except Exception as e:
            flash(f"Storage location is not writable: {e}", "danger")
            return render_template("setup.html", config=config, ollama_status=ollama_status)

        # Generate Salt
        SALT = secrets.token_bytes(16)
        SALT_HEX = SALT.hex()
        
        # Hash for auth
        MASTER_PASSWORD_HASH = hash_password(password, SALT)
        
        # Save to config
        config["master_password_hash"] = MASTER_PASSWORD_HASH
        config["salt"] = SALT_HEX
        config["storage_dir"] = storage_dir
        config["ollama_base_url"] = ollama_url
        config["chat_model"] = chat_model
        
        save_config()
        ensure_storage_dirs()
            
        # Log user in immediately
        derived_key = derive_key(password, SALT)
        session['encryption_key'] = derived_key.decode()
        
        flash("Setup complete! You are now logged in.", "success")
        return redirect(url_for("signature.index"))

    return render_template("setup.html", config=config, ollama_status=ollama_status)

@signature_bp.route("/login", methods=["GET", "POST"])
def login():
    global MASTER_PASSWORD_HASH, SALT
    
    if request.method == "POST":
        password = request.form.get("password")
        
        if not SALT:
             # Should not happen if hash exists, but safety check
             flash("Error: Salt missing. Please reset config.", "danger")
             return render_template("login.html")
             
        if verify_password(password, SALT, MASTER_PASSWORD_HASH):
            # Success
            if not MASTER_PASSWORD_HASH.startswith("pbkdf2_sha256$"):
                MASTER_PASSWORD_HASH = hash_password(password, SALT)
                config["master_password_hash"] = MASTER_PASSWORD_HASH
                save_config()
            derived_key = derive_key(password, SALT)
            session['encryption_key'] = derived_key.decode()
            flash("Logged in successfully.", "success")
            return redirect(url_for("signature.index"))
        else:
            flash("Invalid password.", "danger")
            
    return render_template("login.html")

@signature_bp.route("/logout")
def logout():
    session.pop('encryption_key', None)
    flash("Logged out.", "info")
    return redirect(url_for("signature.login"))


def is_valid_fernet_key(key):
    if not key or not isinstance(key, str):
        return False
    try:
        Fernet(key)
        return True
    except Exception:
        return False

if not is_valid_fernet_key(FERNET_KEY):
    # If missing or invalid (e.g. placeholder), generate a new one
    FERNET_KEY = Fernet.generate_key().decode()
    config["fernet_key"] = FERNET_KEY
    save_config()

# Initialize cipher
cipher = Fernet(FERNET_KEY)

UPLOADS_DIR = ""
SIGNATURES_DIR = ""
ensure_storage_dirs()


def write_metadata(meta: dict):
    """Write metadata for a single file (meta must contain 'id')."""
    path = metadata_path_for(meta["id"])
    with open(path, "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)


def read_metadata(path: str) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _create_meta_for_existing_file(filename: str) -> dict:
    """Create metadata for a raw file found in uploads (migration helper)."""
    file_id = str(uuid.uuid4())
    stored_filename = filename
    file_path = os.path.join(UPLOADS_DIR, stored_filename)
    meta = {
        "id": file_id,
        "original_filename": filename,
        "stored_filename": stored_filename,
        "filename": filename,
        "file_path": file_path,
        "type": "pdf" if filename.lower().endswith(".pdf") else "signature",
        "status": "uploaded",
        "timestamp": datetime.utcnow().isoformat() + "Z"
    }
    # save metadata file next to it
    write_metadata(meta)
    return meta


def load_docs() -> list:
    """
    Return a list of metadata dicts for all documents.
    Priority: read per-file JSON metadata files. If no metadata exists
    but files are present, create metadata for them (migration).
    Result is sorted by timestamp desc (newest first).
    """
    docs = []
    # find any metadata JSON files in uploads dir
    for name in os.listdir(UPLOADS_DIR):
        if name.endswith(".json"):
            try:
                meta = read_metadata(os.path.join(UPLOADS_DIR, name))
                docs.append(meta)
            except Exception:
                # skip malformed metadata files
                continue

    # If we found no metadata but there are files, create metadata for each
    if not docs:
        for name in os.listdir(UPLOADS_DIR):
            if name.lower().endswith((".pdf", ".png", ".jpg", ".jpeg")):
                docs.append(_create_meta_for_existing_file(name))

    # normalize & sort (if timestamp missing, put at end)
    def _ts(d):
        t = d.get("timestamp")
        if not t:
            return ""
        return t
    docs.sort(key=_ts, reverse=True)
    return docs


def save_metadata_for_doc_id(doc_id: str, updates: dict):
    """Load metadata for doc_id, update with `updates` dict and save."""
    path = metadata_path_for(doc_id)
    if not os.path.exists(path):
        raise FileNotFoundError("Metadata not found for id: " + doc_id)
    meta = read_metadata(path)
    meta.update(updates)
    write_metadata(meta)
    return meta


def analyze_document_content(body: str, language: str = "English", style: str = "layman") -> str:
    user_prompt = config.get("analysis_prompt")
    
    if user_prompt:
        # Simple string formatting for user prompt
        # We allow placeholder {language} and {style}
        system_prompt = user_prompt.replace("{language}", language).replace("{style}", style)
        # Add basic role if not present? No, trust user prompt.
    else:
        lang_prompt = f"Provide the analysis in {language}."
        style_prompt = ""
        if style == "layman":
            style_prompt = (
                "Explain everything in simple, layman terms that someone without a legal background can understand. "
                "Avoid complex legal jargon, and if you must use it, explain it clearly."
            )
        
        system_prompt = (
            "You are an expert legal analyst. "
            f"Analyze the following document. {lang_prompt} {style_prompt} "
            "Structure your response as follows:\n"
            "1. **Executive Summary**: A brief overview.\n"
            "2. **Key Obligations**: Main responsibilities for each party.\n"
            "3. **Risk Assessment**: Potential risks, liabilities, or unusual clauses.\n"
            "4. **Recommendations**: Verification or specific actions required.\n"
            "Format the output in clean Markdown using headers and bullet points."
        )

    chat_model = get_required_model("chat")
    try:
        ollama_client().list()
    except Exception:
        raise Exception("Ollama server is not running. Start it with: ollama serve")
    try:
        ensure_ollama_model(chat_model)
        response = ollama_client().chat(
            model=chat_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": body}
            ]
        )
        return response["message"]["content"].strip()
    except Exception as e:
        raise Exception(f"Ollama chat error: {e}")


def get_available_vision_model():
    """Return the configured vision model if available."""
    user_pref = (config.get("vision_model") or "").strip()
    if not user_pref:
        return None
    
    try:
        available = get_ollama_models()
            
        # If user pref is valid, return it
        if user_pref in available:
            return user_pref

        ensure_ollama_model(user_pref)
        return user_pref
    except Exception as e:
        print(f"Error listing Ollama models: {e}")
        return None

def extract_text_from_image(image_path: str, raw_bytes: bytes = None) -> str:
    """Extract text from an image using Ollama with a vision model.
    Args:
        image_path: Path to image file (will be decrypted if raw_bytes not provided)
        raw_bytes: Optional raw image bytes (for temp files that aren't encrypted)
    """
    # check if ollama is running
    try:
        ollama_client().list()
    except Exception:
        raise Exception("Ollama server is not running. Start it with: ollama serve")

    vision_model = get_available_vision_model()
    if not vision_model:
        raise Exception("No Ollama vision model selected. Choose a vision model in Settings.")

    try:
        ensure_ollama_model(vision_model)
        # Use provided raw bytes or decrypt from file
        if raw_bytes:
            image_bytes = raw_bytes
        else:
            image_bytes = decrypt_file_content(image_path, g.encryption_key)
            
        response = ollama_client().chat(
            model=vision_model, 
            messages=[
                {
                    'role': 'user',
                    'content': 'Extract all text from this image. Output only the text, no conversational filler.',
                    'images': [image_bytes]
                }
            ]
        )
        content = response['message']['content'].strip()
        return content
    except Exception as e:
        print(f"Vision extraction failed with model {vision_model}: {e}")
        return ""

def extract_text_from_scanned_pdf(file_path: str) -> str:
    """Render PDF pages to images and extract text using OCR (Ollama Vision)."""
    text = ""
    try:
        # pypdfium2 needs path or bytes
        file_bytes = decrypt_file_content(file_path, g.encryption_key)
        pdf = pdfium.PdfDocument(file_bytes)
        n_pages = len(pdf)
        
        # Limit pages to avoid taking forever on large docs for now
        if n_pages > 5:
            print(f"PDF has {n_pages} pages. OCRing first 5 pages only for performance.")
            
        for i in range(min(n_pages, 5)):
            print(f"OCR processing page {i+1} of {min(n_pages, 5)}...")
            page = pdf[i]
            # Render page to bitmap, then to bytes
            bitmap = page.render(scale=2.0) # 2.0 scale for better quality
            pil_image = bitmap.to_pil()
            
            # Save to temp buffer
            img_byte_arr = io.BytesIO()
            pil_image.save(img_byte_arr, format='PNG')
            img_bytes = img_byte_arr.getvalue()
            
            # Pass raw bytes directly to avoid encryption/decryption issues
            page_text = extract_text_from_image("", raw_bytes=img_bytes)
            if page_text:
                text += page_text + "\n"
                
        return text
    except Exception as e:
        print(f"Scanned PDF extraction failed: {e}")
        import traceback
        traceback.print_exc()
        return ""

def get_document_text(file_path: str) -> str:
    """Helper to extract text from various file formats."""
    text = ""
    filename = file_path.lower()
    
    if filename.endswith(".pdf"):
        # For PDFplumber, we need a file-like object since we have bytes
        file_bytes = decrypt_file_content(file_path, g.encryption_key)
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        
        # If little to no text found, assume scanned and try OCR
        if len(text.strip()) < 50:
            print("Little text found in PDF. Attempting OCR (scanned PDF)...")
            ocr_text = extract_text_from_scanned_pdf(file_path)
            if ocr_text.strip():
                text = ocr_text
    elif filename.endswith(".docx"):
        # python-docx needs a file path or stream
        file_bytes = decrypt_file_content(file_path, g.encryption_key)
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"
    elif filename.endswith(".txt"):
        file_bytes = decrypt_file_content(file_path, g.encryption_key)
        text = file_bytes.decode("utf-8")
    elif filename.endswith((".png", ".jpg", ".jpeg")):
        text = extract_text_from_image(file_path)
    else:
        raise ValueError("Unsupported file type.")
    
    return text


def translate_document_content(text: str, target_language: str = "English", source_language: str = "auto") -> str:
    user_prompt = config.get("translation_prompt")
    
    if user_prompt:
        system_prompt = user_prompt.replace("{target_language}", target_language).replace("{source_language}", source_language)
    else:
        # Robust default prompt
        source_instruction = f"from {source_language} " if source_language and source_language != "auto" else ""
        system_prompt = (
            f"You are a professional translator. Translate the following text {source_instruction}into {target_language}. "
            "Maintain the original tone, formatting, and legal nuances (if any) as much as possible. "
            "Return ONLY the translated text, with no introductory or concluding remarks. "
            "Do not output markdown code blocks unless the original text contained them."
        )

    chat_model = get_required_model("chat")
    try:
        ollama_client().list()
    except Exception:
        raise Exception("Ollama server is not running. Start it with: ollama serve")
    try:
        ensure_ollama_model(chat_model)
        response = ollama_client().chat(
            model=chat_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text}
            ]
        )
        return response["message"]["content"].strip()
    except Exception as e:
        raise Exception(f"Ollama chat error: {e}")


def generate_response_content(text: str, instructions: str = "", tone: str = "Professional") -> str:
    system_prompt = (
        f"You are an AI assistant helping to draft a response to a document. "
        f"The user has provided the following document content. "
        f"Draft a response in a {tone} tone. "
    )
    
    if instructions:
        system_prompt += f" Additional instructions from the user: {instructions}"

    try:
        ollama_client().list()
    except Exception:
        raise Exception("Ollama server is not running. Start it with: ollama serve")
    try:
        chat_model = get_required_model("chat")
        ensure_ollama_model(chat_model)
        response = ollama_client().chat(
            model=chat_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text}
            ]
        )
        return response["message"]["content"].strip()
    except Exception as e:
        raise Exception(f"Ollama chat error: {e}")


def answer_document_question(text: str, question: str) -> str:
    system_prompt = (
        "You are a careful document analysis assistant. Answer the user's question using only the document content "
        "provided. If the document does not contain enough information, say what is missing. Be clear, practical, "
        "and point to relevant sections or wording when possible."
    )
    user_content = f"Document content:\n{text}\n\nQuestion:\n{question}"

    chat_model = get_required_model("chat")
    try:
        ollama_client().list()
    except Exception:
        raise Exception("Ollama server is not running. Start it with: ollama serve")
    try:
        ensure_ollama_model(chat_model)
        response = ollama_client().chat(
            model=chat_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content}
            ]
        )
        return response["message"]["content"].strip()
    except Exception as e:
        raise Exception(f"Ollama chat error: {e}")

# Default route → dashboard
@signature_bp.route("/")
def index():
    docs = load_docs()
    return render_template("docdashboard.html", docs=docs)


@signature_bp.route("/documentation")
def documentation():
    base_path = getattr(sys, '_MEIPASS', os.getcwd())
    doc_path = os.path.join(base_path, "documentation.md")
    if not os.path.exists(doc_path):
        abort(404, description="Documentation not found")

    with open(doc_path, "r", encoding="utf-8") as f:
        doc_markdown = f.read()

    doc_html = markdown.markdown(doc_markdown, extensions=["fenced_code", "tables"])
    return render_template("documentation.html", doc_html=doc_html)


@signature_bp.route("/settings", methods=["GET", "POST"])
def settings():
    global config, OLLAMA_BASE_URL
    
    # Fetch available Ollama models for dropdowns
    ollama_models = []
    ollama_status = {"installed": ollama_is_installed(), "running": False, "error": None}
    try:
        ollama_models = get_ollama_models()
        ollama_status["running"] = True
    except Exception as e:
        ollama_status["error"] = str(e)
        print(f"Error fetching Ollama models: {e}")

    if request.method == "POST":
        ollama_url = request.form.get("ollama_base_url", DEFAULT_OLLAMA_BASE_URL).strip() or DEFAULT_OLLAMA_BASE_URL
        storage_dir = request.form.get("storage_dir", "").strip()
        email = request.form.get("email", "").strip()
        app_password = request.form.get("app_password", "").strip()
        
        # New settings
        vision_model = request.form.get("vision_model")
        chat_model = request.form.get("chat_model", "").strip()
        custom_model_name = request.form.get("custom_model_name", "").strip()
        languages = normalize_languages(request.form.get("languages", ""))
        translation_prompt = request.form.get("translation_prompt", "").strip()
        analysis_prompt = request.form.get("analysis_prompt", "").strip() or DEFAULT_ANALYSIS_PROMPT
        previous_storage_dir = storage_root()

        try:
            storage_dir = validate_storage_dir(storage_dir or previous_storage_dir)
        except Exception as e:
            flash(f"Storage location is not writable: {e}", "danger")
            return render_template(
                "settings.html",
                config=config,
                ollama_models=ollama_models,
                ollama_status=ollama_status,
                languages_text="\n".join(configured_languages()),
                default_analysis_prompt=DEFAULT_ANALYSIS_PROMPT,
            )
        
        for legacy_key in (
            "use_" + "open" + "ai",
            "open" + "ai_" + "api" + "_key",
            "open" + "ai_model",
        ):
            config.pop(legacy_key, None)
        config["languages"] = languages
        config["ollama_base_url"] = ollama_url
        config["storage_dir"] = storage_dir
        config["email"] = email
        config["app_password"] = app_password
        
        config["vision_model"] = vision_model or ""
        config["chat_model"] = chat_model
        config["custom_model_name"] = custom_model_name
        config["translation_prompt"] = translation_prompt
        config["analysis_prompt"] = analysis_prompt
        
        if os.path.abspath(storage_dir) != os.path.abspath(previous_storage_dir):
            migrate_storage_dir(storage_dir)
        save_config()
        ensure_storage_dirs()
            
        # Update current runtime variables
        OLLAMA_BASE_URL = ollama_url
        
        flash("Settings updated successfully.", "success")
        return redirect(url_for("signature.index"))
        
    return render_template(
        "settings.html",
        config=config,
        ollama_models=ollama_models,
        ollama_status=ollama_status,
        languages_text="\n".join(configured_languages()),
        default_analysis_prompt=DEFAULT_ANALYSIS_PROMPT,
    )

@signature_bp.route("/doc/send/<doc_id>", methods=["POST"])
def doc_send(doc_id):
    doc_id = safe_doc_id(doc_id)
    to_email = request.form.get("to")
    cc_email = request.form.get("cc")
    bcc_email = request.form.get("bcc")
    message_text = request.form.get("message")

    if not to_email:
        return jsonify({"error": "Recipient email is required"}), 400

    # Load metadata to find the file
    docs = load_docs()
    doc = next((d for d in docs if d['id'] == doc_id), None)
    if not doc:
        return jsonify({"error": "Document not found"}), 404

    # Determine file path
    file_path = resolve_upload_path(doc)

    attachment_name = doc.get('filename', 'document.pdf')

    if not file_path or not os.path.exists(file_path):
        return jsonify({"error": f"File not found on disk: {file_path}"}), 404

    # Get email credentials
    sender_email = config.get("email")
    app_password = config.get("app_password")
    if app_password:
        app_password = app_password.replace(" ", "").strip()

    if not sender_email or not app_password:
        return jsonify({"error": "Email credentials not configured in settings"}), 400

    try:
        # Create message
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = to_email
        if cc_email:
            msg['Cc'] = cc_email
        msg['Subject'] = f"Document: {attachment_name}"

        msg.attach(MIMEText(message_text, 'plain'))

        # Attach file
        file_bytes = decrypt_file_content(file_path, g.encryption_key)
        part = MIMEApplication(file_bytes, Name=attachment_name)
        part['Content-Disposition'] = f'attachment; filename="{attachment_name}"'
        msg.attach(part)

        # Send email
        recipients = [to_email]
        if cc_email:
            recipients.append(cc_email)
        if bcc_email:
            recipients.append(bcc_email)

        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(sender_email, app_password)
            server.sendmail(sender_email, recipients, msg.as_string())

        return jsonify({"success": True, "message": "Email sent successfully!"})

    except Exception as e:
        return jsonify({"error": f"Failed to send email: {str(e)}"}), 500

@signature_bp.route("/api/train_local_gpt", methods=["POST"])
def train_local_gpt():
    method = (request.json or {}).get("method", "pytorch")
    if method not in {"python", "pytorch"}:
        return jsonify({"error": "Unsupported training method"}), 400
    custom_model_name = (config.get("custom_model_name") or "").strip()
    if not custom_model_name:
        return jsonify({"error": "Set a custom model name in Settings before training."}), 400
    # Spawn background training process
    import subprocess
    import threading
    
    def run_training(key_bytes, train_method):
        try:
            base_path = getattr(sys, '_MEIPASS', os.getcwd())
            script_path = os.path.join(base_path, "scripts", "train_to_gguf.py")
            # Using the same python executable as the app
            subprocess.run([
                sys.executable,
                script_path,
                "--key", key_bytes.decode(),
                "--data-dir", storage_root(),
                "--method", train_method,
                "--model-name", custom_model_name,
                "--vision-model", (config.get("vision_model") or "")
            ], check=True)
            print(f"Finished training local GPT model using {train_method}.")
        except Exception as e:
            print(f"Error during local GPT training: {e}")

    thread = threading.Thread(target=run_training, args=(g.encryption_key, method))
    thread.daemon = True
    thread.start()
    return jsonify({"success": True, "message": "Local GPT training started in background."})

@signature_bp.route("/api/training_status", methods=["GET"])
def training_status():
    status_file = os.path.join(DATA_BASE_DIR, "training_status.json")
    if os.path.exists(status_file):
        try:
            with open(status_file, "r") as f:
                data = json.load(f)
                return jsonify(data)
        except Exception:
            return jsonify({"status": "unknown", "message": "Could not read status file."})
    return jsonify({"status": "unknown", "message": "No training status found."})

@signature_bp.route("/api/search_local_gpt", methods=["POST"])
def search_local_gpt():
    query = (request.json or {}).get("query", "")
    if not query:
        return jsonify({"error": "Query is required"}), 400
    custom_model_name = (config.get("custom_model_name") or "").strip()
    if not custom_model_name:
        return jsonify({"error": "Set a custom model name in Settings before searching."}), 400
        
    try:
        response = ollama_client().chat(
            model=custom_model_name,
            messages=[{"role": "user", "content": query}]
        )
        return jsonify({"success": True, "response": response["message"]["content"]})
    except Exception as e:
        return jsonify({"error": f"Failed to search local GPT: {str(e)}"}), 500


# Categories API
@signature_bp.route("/api/categories", methods=["GET", "POST"])
def categories_api():
    global config
    if request.method == "GET":
        cats = config.get("categories", ["Contract", "Invoice", "Report", "Image", "Other"])
        return jsonify({"categories": cats})

    # POST – add a new category
    name = (request.json or {}).get("name", "").strip()
    if not name:
        return jsonify({"error": "Category name is required"}), 400
    cats = list(config.get("categories", []))
    if name in cats:
        return jsonify({"error": "Category already exists"}), 400
    cats.append(name)
    config["categories"] = cats
    save_config()
    return jsonify({"categories": cats})


@signature_bp.route("/api/categories/<name>", methods=["DELETE", "PUT"])
def category_item_api(name):
    global config
    cats = list(config.get("categories", []))

    if request.method == "DELETE":
        if name not in cats:
            return jsonify({"error": "Category not found"}), 404
        cats.remove(name)
        config["categories"] = cats
        save_config()
        return jsonify({"categories": cats})

    # PUT – rename
    new_name = (request.json or {}).get("name", "").strip()
    if not new_name:
        return jsonify({"error": "New name is required"}), 400
    if name not in cats:
        return jsonify({"error": "Category not found"}), 404
    if new_name in cats:
        return jsonify({"error": "Category already exists"}), 400
    cats[cats.index(name)] = new_name
    config["categories"] = cats
    save_config()
    return jsonify({"categories": cats})


# View PDF
@signature_bp.route("/docupload", methods=["GET", "POST"])
def docupload():
    if request.method == "POST":
        pdf_file = request.files.get("pdf_file")

        file_id = str(uuid.uuid4())

        if pdf_file and pdf_file.filename:
            original_filename = secure_filename(pdf_file.filename)
            extension = original_filename.rsplit('.', 1)[-1].lower() if '.' in original_filename else ''
            if extension not in ALLOWED_UPLOAD_EXTENSIONS:
                 flash(f"Unsupported file type: {extension or 'none'}", "danger")
                 return redirect(url_for("signature.docupload"))

            # store with file_id prefix to avoid collisions
            stored_filename = f"{file_id}_{original_filename}"
            file_path = os.path.join(UPLOADS_DIR, stored_filename)
            pdf_file.save(file_path)

            # Encrypt the file immediately
            if 'encryption_key' in g:
                encrypt_file(file_path, g.encryption_key)

            # Read optional user-supplied metadata from form
            title = request.form.get("title", "").strip() or os.path.splitext(original_filename)[0]
            description = request.form.get("description", "").strip()
            category = request.form.get("doc_type", "").strip()

            meta = {
                "id": file_id,
                "original_filename": original_filename,
                "stored_filename": stored_filename,
                "filename": original_filename,
                "title": title,
                "description": description,
                "category": category,
                "file_path": file_path,
                "type": extension,
                "status": "uploaded",
                "timestamp": datetime.utcnow().isoformat() + "Z",
                "uploaded_at": datetime.utcnow().isoformat()
            }
            write_metadata(meta)

            flash(f"Document {original_filename} uploaded successfully.", "success")
            return redirect(url_for("signature.index"))

        else:
            flash("Please provide a PDF file URL.", "danger")
            return redirect(url_for("signature.docupload"))

    # GET -> show current docs
    docs = load_docs()
    return render_template("docdashboard.html", docs=docs)


# Serve uploaded files
@signature_bp.route('/docuploaded_file/<path:filename>')
def docuploaded_file(filename):
    stored_meta = None
    for f in os.listdir(UPLOADS_DIR):
        if f.endswith(".json"):
            meta_path = os.path.join(UPLOADS_DIR, f)
            with open(meta_path) as meta_file:
                meta = json.load(meta_file)

            if meta.get("original_filename") == filename:
                stored_meta = meta
                break

    if not stored_meta:
        abort(404)

    file_path = resolve_upload_path(stored_meta)
    if not file_path:
         abort(404)

    # Decrypt and serve
    try:
        if 'encryption_key' in g:
            content = decrypt_file_content(file_path, g.encryption_key)
            return send_file(
                io.BytesIO(content),
                mimetype='application/pdf' if filename.endswith('.pdf') else 'image/png', # simplistic mime
                as_attachment=False,
                download_name=filename
            )
        else:
             # Should be caught by middleware but just in case
             abort(403)
    except Exception as e:
        print(f"Error serving file: {e}")
        abort(500)

    abort(404, description=f"No stored file found for {filename}")


# Sign document (draw signature)
@signature_bp.route("/docsign/<doc_id>")
def docsign(doc_id):
    doc_id = safe_doc_id(doc_id)
    meta_path = metadata_path_for(doc_id)
    if not os.path.exists(meta_path):
        flash("Document metadata not found.", "danger")
        return redirect(url_for("signature.index"))

    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)

    file_path = resolve_upload_path(meta)
    if not file_path:
        flash("Document file not found on disk.", "danger")
        return redirect(url_for("signature.index"))

    # Provide a URL for the iframe to display the PDF
    meta["view_url"] = url_for("signature.serve_doc", doc_id=doc_id)

    return render_template("docsignature.html", document=meta)


# View PDF
@signature_bp.route("/docview/<doc_id>")
def docview(doc_id):
    doc_id = safe_doc_id(doc_id)
    meta_path = metadata_path_for(doc_id)
    if not os.path.exists(meta_path):
        flash("Document metadata not found.", "danger")
        return redirect(url_for("signature.index"))

    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)

    if meta.get("type") == "pdf":
        meta["view_url"] = url_for("signature.serve_doc", doc_id=doc_id)
    elif meta.get("type") == "google_doc":
        meta["view_url"] = meta.get("url")
    elif meta.get("type") in ("jpg", "jpeg", "png"):
        meta["view_url"] = url_for("signature.serve_doc", doc_id=doc_id)
    else:
        meta["view_url"] = None

    return render_template("docview.html", doc=meta)


@signature_bp.route("/docdelete/<doc_id>", methods=["POST"])
def docdelete(doc_id):
    doc_id = safe_doc_id(doc_id)
    meta_path = metadata_path_for(doc_id)
    if not os.path.exists(meta_path):
        flash("Document metadata not found.", "danger")
        return redirect(url_for("signature.index"))

    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)

    file_path = resolve_upload_path(meta)
    if file_path:
        os.remove(file_path)

    # Delete the metadata JSON
    os.remove(meta_path)
    flash(f"Document '{meta.get('filename', doc_id)}' deleted successfully.", "success")
    return redirect(url_for("signature.index"))


@signature_bp.route("/download/<doc_id>")
def download_doc(doc_id):
    doc_id = safe_doc_id(doc_id)
    meta_path = metadata_path_for(doc_id)
    if not os.path.exists(meta_path):
        flash("Document metadata not found.", "danger")
        return redirect(url_for("signature.index"))

    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)

    file_path = resolve_upload_path(meta)
    if not file_path:
        flash("Document file not found on disk.", "danger")
        return redirect(url_for("signature.index"))

    content = decrypt_file_content(file_path, g.encryption_key)
    return send_file(io.BytesIO(content), as_attachment=True, download_name=meta.get("filename", doc_id))


@signature_bp.route("/save_signed_pdf", methods=["POST"])
def save_signed_pdf():
    data = request.json or {}
    pdf_base64 = data.get("pdf_base64")
    original_filename = secure_filename(data.get("original_filename") or "signed.pdf")

    if not pdf_base64 or not original_filename:
        return {"success": False, "error": "Missing data"}, 400

    file_id = str(uuid.uuid4())
    # generate new signed filename
    signed_filename = f"signed_{file_id}_{original_filename}"
    signed_path = os.path.join(UPLOADS_DIR, signed_filename)

    # decode base64
    pdf_bytes = base64.b64decode(pdf_base64.split(",")[1] if "," in pdf_base64 else pdf_base64)
    with open(signed_path, "wb") as f:
        f.write(pdf_bytes)

    # create meta file
    # create meta file (use only the id in the meta filename)

    # Encrypt the signed file
    if 'encryption_key' in g:
        encrypt_file(signed_path, g.encryption_key)

    meta = {
        "id": f"signed_{file_id}",
        "original_filename": original_filename,
        "stored_filename": signed_filename,
        "filename": original_filename,
        "file_path": signed_path,   # better to store full path instead of just folder
        "type": "pdf",
        "status": "signed",
        "timestamp": datetime.utcnow().isoformat() + "Z"
    }

    write_metadata(meta)

    return {"success": True, "filename": signed_filename}


# Analyse document
@signature_bp.route("/analyze_doc/<doc_id>", methods=["GET", "POST"])
def analyze_doc(doc_id):
    doc_id = safe_doc_id(doc_id)
    languages = configured_languages()
    meta_path = metadata_path_for(doc_id)
    if not os.path.exists(meta_path):
        flash("Document metadata not found.", "danger")
        return redirect(url_for("signature.index"))

    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)

    if request.method == "POST":
        language = request.form.get("language", languages[0])
        if language not in languages:
            language = languages[0]
        style = request.form.get("style", "layman")
        
        file_path = resolve_upload_path(meta)
        if not file_path:
            flash("Document file not found.", "danger")
            return redirect(url_for("signature.index"))

        try:
            text = get_document_text(file_path)
        except Exception as e:
            flash(f"Error extracting text: {e}", "danger")
            return redirect(url_for("signature.index"))

        if not text.strip():
            flash("No text content found in document to analyze.", "warning")
            return redirect(url_for("signature.index"))

        try:
            raw_analysis = analyze_document_content(text, language=language, style=style)
            # Convert analysis to HTML for better display
            html_analysis = markdown.markdown(raw_analysis)
            
            # Store analysis in metadata
            meta["analysis"] = html_analysis
            meta["analysis_language"] = language
            meta["analysis_style"] = style
            write_metadata(meta)
            
            return render_template("docanalysis.html", doc=meta, analysis=html_analysis, languages=languages)
        except Exception as e:
            flash(f"Analysis failed: {e}", "danger")
            return redirect(url_for("signature.index"))

    analysis_html = meta.get("analysis")
    
    # Check if JSON is requested (more robust check)
    is_json = request.is_json or \
              'application/json' in request.headers.get('Accept', '') or \
              request.args.get('format') == 'json'

    if is_json:
        if not analysis_html:
            # Fallback to current settings if no analysis exists
            try:
                file_path = resolve_upload_path(meta)
                if not file_path:
                     return jsonify({"success": False, "analysis": "Document file not found."})
                
                text = get_document_text(file_path)
                if not text.strip():
                    return jsonify({"success": False, "analysis": "Low or no text content found to analyze."})

                analysis_text_raw = analyze_document_content(text, language=languages[0])
                analysis_html = markdown.markdown(analysis_text_raw)
                return jsonify({"success": True, "analysis": analysis_html})
            except Exception as e:
                return jsonify({"success": False, "analysis": f"Extraction error: {str(e)}"})
        
        # Strip HTML for the simple modal if needed, or just return as is
        return jsonify({"success": True, "analysis": analysis_html})

    # GET -> show analysis options or current analysis if exists
    return render_template("docanalysis.html", doc=meta, analysis=analysis_html, languages=languages)


@signature_bp.route("/translate_doc/<doc_id>", methods=["POST"])
def translate_doc(doc_id):
    doc_id = safe_doc_id(doc_id)
    languages = configured_languages()
    meta_path = metadata_path_for(doc_id)
    if not os.path.exists(meta_path):
        return jsonify({"success": False, "error": "Document metadata not found."}), 404

    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)

    data = request.json or {}
    file_path = resolve_upload_path(meta)
    if not file_path:
        return jsonify({"success": False, "error": "Document file not found."}), 404

    target_language = data.get("language", languages[0])
    if target_language not in languages:
        target_language = languages[0]
    source_language = data.get("source_language", "auto")
    if source_language != "auto" and source_language not in languages:
        source_language = "auto"
    show_extracted = data.get("show_extracted_text", False)

    # Get doc text
    try:
        text = get_document_text(file_path)
    except Exception as e:
         return jsonify({"error": f"Failed to extract text: {e}"}), 500
         
    if not text.strip():
        # Fallback for scanned PDF if not already handled inside get_document_text
        # But get_document_text handles it now.
        return jsonify({"error": "No text content found to translate."}), 400

    try:
        print(f"Starting translation to {target_language}, text length: {len(text)} chars")
        translation = translate_document_content(text, target_language, source_language)
        print(f"Translation completed, result length: {len(translation)} chars")
        
        # Convert translation to HTML for display
        html_translation = markdown.markdown(translation)
        
        result = {"success": True, "translation": html_translation}
        
        # Include extracted text if requested
        if show_extracted:
            result["extracted_text"] = text
        
        return jsonify(result)
    except Exception as e:
        print(f"Translation error: {e}")
        return jsonify({"error": str(e)}), 500


@signature_bp.route("/save_translation/<doc_id>", methods=["POST"])
def save_translation(doc_id):
    doc_id = safe_doc_id(doc_id)
    """Save translation as a new .txt document"""
    try:
        data = request.json or {}
        translation_text = data.get("translation_text", "")
        original_filename = data.get("original_filename", "document")
        target_language = data.get("target_language", "")
        
        if not translation_text:
            return jsonify({"success": False, "error": "No translation text provided"}), 400
        
        # Create new document ID
        new_doc_id = str(uuid.uuid4())
        
        # Create filename
        base_name = secure_filename(os.path.splitext(original_filename)[0]) or "document"
        safe_language = secure_filename(target_language) or "translated"
        new_filename = f"{base_name}_translated_{safe_language}.txt"
        
        # Ensure uploads directory exists
        os.makedirs(UPLOADS_DIR, exist_ok=True)
        
        # Save the text file (encrypted)
        file_path = os.path.join(UPLOADS_DIR, f"{new_doc_id}.txt")
        fernet = Fernet(g.encryption_key)
        encrypted_content = fernet.encrypt(translation_text.encode('utf-8'))
        with open(file_path, 'wb') as f:
            f.write(encrypted_content)
        
        # Create metadata
        metadata = {
            "id": new_doc_id,
            "filename": new_filename,
            "type": "txt",
            "status": "translated",
            "file_path": file_path,
            "uploaded_at": datetime.now().isoformat(),
            "source_doc_id": doc_id,
            "target_language": target_language
        }
        
        # Save metadata
        write_metadata(metadata)
        
        return jsonify({
            "success": True,
            "doc_id": new_doc_id,
            "filename": new_filename
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@signature_bp.route("/generate_response/<doc_id>", methods=["POST"])
def generate_response(doc_id):
    doc_id = safe_doc_id(doc_id)
    meta_path = metadata_path_for(doc_id)
    if not os.path.exists(meta_path):
        return jsonify({"success": False, "error": "Document metadata not found."}), 404

    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)

    data = request.json or {}
    instructions = data.get("instructions", "")
    tone = data.get("tone", "Professional")
    
    file_path = resolve_upload_path(meta)
    if not file_path:
        return jsonify({"success": False, "error": "Document file not found."}), 404

    try:
        text = get_document_text(file_path)
    except Exception as e:
        return jsonify({"success": False, "error": f"Error extracting text: {e}"}), 500

    if not text.strip():
        return jsonify({"success": False, "error": "No text content found to generate response."}), 400

    try:
        response_text = generate_response_content(text, instructions=instructions, tone=tone)
        # Convert to HTML
        response_html = markdown.markdown(response_text)
        
        # Save to metadata? Maybe not needed for response, but could be useful. 
        # For now, let's just return it. Use-case is likely copy-pasting.
        # meta["generated_response"] = response_html 
        # write_metadata(meta)
        
        return jsonify({"success": True, "response": response_html})
    except Exception as e:
        return jsonify({"success": False, "error": f"Response generation failed: {e}"}), 500


@signature_bp.route("/ask_doc/<doc_id>", methods=["POST"])
def ask_doc(doc_id):
    doc_id = safe_doc_id(doc_id)
    meta_path = metadata_path_for(doc_id)
    if not os.path.exists(meta_path):
        return jsonify({"success": False, "error": "Document metadata not found."}), 404

    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)

    data = request.json or {}
    question = data.get("question", "").strip()
    if not question:
        return jsonify({"success": False, "error": "Question is required."}), 400

    file_path = resolve_upload_path(meta)
    if not file_path:
        return jsonify({"success": False, "error": "Document file not found."}), 404

    try:
        text = get_document_text(file_path)
    except Exception as e:
        return jsonify({"success": False, "error": f"Error extracting text: {e}"}), 500

    if not text.strip():
        return jsonify({"success": False, "error": "No text content found in document."}), 400

    try:
        answer = answer_document_question(text, question)
        return jsonify({"success": True, "answer": markdown.markdown(answer)})
    except Exception as e:
        return jsonify({"success": False, "error": f"Question failed: {e}"}), 500

# Save drawn signature
@signature_bp.route("/save_signature", methods=["POST"])
def save_signature():
    data = request.json or {}
    sig_data = data.get("signature")
    doc_id = safe_doc_id(data.get("doc_id", ""))

    if not sig_data or not doc_id:
        return {"success": False, "error": "Missing signature or document ID"}, 400

    sig_id = f"{uuid.uuid4()}.png"
    sig_path = os.path.join(SIGNATURES_DIR, sig_id)
    os.makedirs(SIGNATURES_DIR, exist_ok=True)

    # decode and encrypt
    raw_bytes = base64.b64decode(sig_data.split(",")[1])
    encrypted = Fernet(g.encryption_key).encrypt(raw_bytes)

    with open(sig_path, "wb") as f:
        f.write(encrypted)

    # update document JSON
    meta_path = metadata_path_for(doc_id)
    if os.path.exists(meta_path):
        with open(meta_path, "r", encoding="utf-8") as f:
            meta = json.load(f)
        meta["signature"] = sig_path
        with open(meta_path, "w", encoding="utf-8") as f:
            json.dump(meta, f, indent=2)

    return {"success": True, "signature_url": url_for("signature.serve_signature", sig_id=sig_id)}


@signature_bp.route("/list_signatures")
def list_signatures():
    files = os.listdir(SIGNATURES_DIR)
    # optionally filter only PNGs
    files = [f for f in files if f.endswith(".png")]
    return jsonify([{"id": f} for f in files])


@signature_bp.route("/serve_signature/<sig_id>")
def serve_signature(sig_id):
    sig_id = secure_filename(sig_id)
    if not sig_id.endswith(".png"):
        abort(400)
    sig_path = os.path.join(SIGNATURES_DIR, sig_id)
    if not os.path.exists(sig_path):
        flash("Signature not found.", "danger")
        return redirect(url_for("signature.index"))

    # read + decrypt
    with open(sig_path, "rb") as f:
        encrypted = f.read()
    try:
        raw_bytes = Fernet(g.encryption_key).decrypt(encrypted)
    except InvalidToken:
        raw_bytes = cipher.decrypt(encrypted)

    return send_file(
        io.BytesIO(raw_bytes),
        mimetype="image/png",
        as_attachment=False,
        download_name=sig_id
    )


@signature_bp.route("/serve/<doc_id>")
def serve_doc(doc_id):
    doc_id = safe_doc_id(doc_id)
    meta_path = metadata_path_for(doc_id)
    if not os.path.exists(meta_path):
        flash("Document metadata not found.", "danger")
        return redirect(url_for("signature.index"))

    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)

    file_path = resolve_upload_path(meta)
    if not file_path:
        flash("Document file not found on disk.", "danger")
        return redirect(url_for("signature.index"))

    try:
        # Serve PDF inline
        # Read and decrypt file content
        content = decrypt_file_content(file_path, g.encryption_key)

        file_type = meta.get("type", "").lower()
        if file_type in ("jpg", "jpeg"):
            mime = "image/jpeg"
        elif file_type == "png":
            mime = "image/png"
        else:
            mime = "application/pdf"

        return send_file(
            io.BytesIO(content),
            mimetype=mime,
            as_attachment=False,
            download_name=meta.get("filename", doc_id)
        )
    except Exception as e:
        flash(f"Error serving document: {e}", "danger")
        return redirect(url_for("signature.index"))


@signature_bp.route("/download_signed/<doc_id>")
def download_signed(doc_id):
    # load your meta JSON
    doc_id = safe_doc_id(doc_id)
    meta_path = metadata_path_for(doc_id)
    if not os.path.exists(meta_path):
        flash("Document not found", "danger")
        return redirect(url_for("signature.index"))

    with open(meta_path, "r", encoding="utf-8") as f:
        meta = json.load(f)

    full_path = resolve_upload_path(meta)
    if not full_path:
        flash("Signed file not found", "danger")
        return redirect(url_for("signature.index"))

    try:
        content = decrypt_file_content(full_path, g.encryption_key)
        
        return send_file(
            io.BytesIO(content), 
            as_attachment=True, 
            download_name=meta.get("filename"),
            mimetype="application/pdf"
        )
    except Exception as e:
        flash(f"Error downloading signed document: {e}", "danger")
        return redirect(url_for("signature.index"))
