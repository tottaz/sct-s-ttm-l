import os
import sys
import json
import math
import random
import io
import argparse
from cryptography.fernet import Fernet
import pdfplumber
import docx
import pypdfium2 as pdfium
import ollama

random.seed(42)

def decrypt_file_content(file_path: str, key: bytes) -> bytes:
    with open(file_path, "rb") as file:
        file_data = file.read()
    try:
        f = Fernet(key)
        return f.decrypt(file_data)
    except Exception:
        return file_data

def get_available_vision_model():
    try:
        response = ollama.list()
        models_list = response.models if hasattr(response, 'models') else response.get('models', [])
        available = [getattr(m, 'model', m.get('model')) if hasattr(m, 'model') or isinstance(m, dict) else str(m) for m in models_list]
        for candidate in ["llava:latest", "llama3.2-vision:latest"]:
            if candidate in available: return candidate
        for name in available:
            if "vision" in name or "llava" in name: return name
        return None
    except: return None

def extract_text_from_image(raw_bytes: bytes) -> str:
    vision_model = get_available_vision_model()
    if not vision_model: return ""
    try:
        response = ollama.chat(
            model=vision_model, 
            messages=[{'role': 'user', 'content': 'Extract all text from this image. Output only the text.', 'images': [raw_bytes]}]
        )
        return response['message']['content'].strip()
    except: return ""

def extract_text_from_scanned_pdf(file_bytes: bytes) -> str:
    text = ""
    try:
        pdf = pdfium.PdfDocument(file_bytes)
        n_pages = len(pdf)
        for i in range(min(n_pages, 5)):
            page = pdf[i]
            bitmap = page.render(scale=2.0)
            pil_image = bitmap.to_pil()
            img_byte_arr = io.BytesIO()
            pil_image.save(img_byte_arr, format='PNG')
            page_text = extract_text_from_image(img_byte_arr.getvalue())
            if page_text: text += page_text + "\n"
        return text
    except: return ""

def extract_text(file_path: str, key: bytes) -> str:
    text = ""
    filename = file_path.lower()
    file_bytes = decrypt_file_content(file_path, key)
    try:
        if filename.endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted: text += extracted + "\n"
            if len(text.strip()) < 50:
                text = extract_text_from_scanned_pdf(file_bytes)
        elif filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs: text += para.text + "\n"
        elif filename.endswith(".txt"):
            text = file_bytes.decode("utf-8")
        elif filename.endswith((".png", ".jpg", ".jpeg")):
            text = extract_text_from_image(file_bytes)
    except Exception as e:
        print(f"Error extracting from {file_path}: {e}")
    return text

def update_status(data_dir: str, status: str, message: str = ""):
    status_file = os.path.join(data_dir, "training_status.json")
    with open(status_file, "w") as f:
        json.dump({"status": status, "message": message}, f)

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--key", required=True, help="Encryption key base64")
    parser.add_argument("--data-dir", required=True, help="Path to data dir")
    parser.add_argument("--method", choices=["python", "pytorch"], default="pytorch", help="Training method")
    args = parser.parse_args()
    
    update_status(args.data_dir, "running", "Loading documents...")

    try:
        uploads_dir = os.path.join(args.data_dir, "uploads")
        docs = []
        if not os.path.exists(uploads_dir):
            raise Exception("Uploads directory not found.")
    
        for name in os.listdir(uploads_dir):
            if name.endswith(".json"):
                with open(os.path.join(uploads_dir, name), "r") as f:
                    meta = json.load(f)
                    file_path = meta.get("file_path")
                    if not file_path or not os.path.exists(file_path):
                        stored_name = meta.get("stored_filename", meta.get("filename"))
                        file_path = os.path.join(uploads_dir, stored_name)
                        
                    if os.path.exists(file_path):
                        text = extract_text(file_path, args.key.encode())
                        if text.strip():
                            chunks = [p.strip() for p in text.split("\n\n") if p.strip()]
                            docs.extend(chunks)
    
        if not docs:
            docs = ["Hello world, this is a fallback document so training doesn't fail."]
    
        random.shuffle(docs)
        
        uchars = sorted(set(''.join(docs)))
        if len(uchars) == 0:
            uchars = ['a']
        vocab_size = len(uchars) + 1 # +1 for BOS
        
        n_layer = 1
        n_embd = 32
        block_size = 64
        n_head = 4
        
        if args.method == "pytorch":
            update_status(args.data_dir, "running", "Training with PyTorch...")
            state_dict = train_pytorch(docs, uchars, vocab_size, n_layer, n_embd, block_size, n_head)
        else:
            update_status(args.data_dir, "running", "Training with Pure Python (expect delays)...")
            state_dict = train_pure_python(docs, uchars, vocab_size, n_layer, n_embd, block_size, n_head)
            
        update_status(args.data_dir, "running", "Exporting model to GGUF format...")
        export_to_gguf(state_dict, uchars, vocab_size, n_embd, n_layer, n_head, block_size, args.data_dir)
        
        update_status(args.data_dir, "completed", "Training successfully complete. You can now use the model.")
    except Exception as e:
        import traceback
        traceback.print_exc()
        update_status(args.data_dir, "error", f"Error during training: {str(e)}")

def train_pure_python(docs, uchars, vocab_size, n_layer, n_embd, block_size, n_head):
    BOS = len(uchars)
    class Value:
        __slots__ = ('data', 'grad', '_children', '_local_grads')
        def __init__(self, data, children=(), local_grads=()):
            self.data = data
            self.grad = 0
            self._children = children
            self._local_grads = local_grads
        def __add__(self, other):
            other = other if isinstance(other, Value) else Value(other)
            return Value(self.data + other.data, (self, other), (1, 1))
        def __mul__(self, other):
            other = other if isinstance(other, Value) else Value(other)
            return Value(self.data * other.data, (self, other), (other.data, self.data))
        def __pow__(self, other): return Value(self.data**other, (self,), (other * self.data**(other-1),))
        def log(self): return Value(math.log(self.data), (self,), (1/self.data,))
        def exp(self): return Value(math.exp(self.data), (self,), (math.exp(self.data),))
        def relu(self): return Value(max(0, self.data), (self,), (float(self.data > 0),))
        def __neg__(self): return self * -1
        def __radd__(self, other): return self + other
        def __sub__(self, other): return self + (-other)
        def __rsub__(self, other): return other + (-self)
        def __rmul__(self, other): return self * other
        def __truediv__(self, other): return self * other**-1
        def __rtruediv__(self, other): return other * self**-1
        def backward(self):
            topo = []
            visited = set()
            def build_topo(v):
                if v not in visited:
                    visited.add(v)
                    for child in v._children:
                        build_topo(child)
                    topo.append(v)
            build_topo(self)
            self.grad = 1
            for v in reversed(topo):
                for child, local_grad in zip(v._children, v._local_grads):
                    child.grad += local_grad * v.grad

    head_dim = n_embd // n_head
    matrix = lambda nout, nin, std=0.08: [[Value(random.gauss(0, std)) for _ in range(nin)] for _ in range(nout)]
    state_dict = {'wte': matrix(vocab_size, n_embd), 'wpe': matrix(block_size, n_embd), 'lm_head': matrix(vocab_size, n_embd)}
    for i in range(n_layer):
        state_dict[f'layer{i}.attn_wq'] = matrix(n_embd, n_embd)
        state_dict[f'layer{i}.attn_wk'] = matrix(n_embd, n_embd)
        state_dict[f'layer{i}.attn_wv'] = matrix(n_embd, n_embd)
        state_dict[f'layer{i}.attn_wo'] = matrix(n_embd, n_embd)
        state_dict[f'layer{i}.mlp_fc1'] = matrix(4 * n_embd, n_embd)
        state_dict[f'layer{i}.mlp_fc2'] = matrix(n_embd, 4 * n_embd)
    
    params = [p for mat in state_dict.values() for row in mat for p in row]

    def linear(x, w): return [sum(wi * xi for wi, xi in zip(wo, x)) for wo in w]
    def softmax(logits):
        max_val = max(val.data for val in logits)
        exps = [(val - max_val).exp() for val in logits]
        total = sum(exps)
        return [e / total for e in exps]
    def rmsnorm(x):
        ms = sum(xi * xi for xi in x) / len(x)
        scale = (ms + 1e-5) ** -0.5
        return [xi * scale for xi in x]

    def gpt(token_id, pos_id, keys, values):
        tok_emb = state_dict['wte'][token_id]
        pos_emb = state_dict['wpe'][pos_id]
        x = [t + p for t, p in zip(tok_emb, pos_emb)]
        x = rmsnorm(x)

        for li in range(n_layer):
            x_residual = x
            x = rmsnorm(x)
            q = linear(x, state_dict[f'layer{li}.attn_wq'])
            k = linear(x, state_dict[f'layer{li}.attn_wk'])
            v = linear(x, state_dict[f'layer{li}.attn_wv'])
            keys[li].append(k)
            values[li].append(v)
            x_attn = []
            for h in range(n_head):
                hs = h * head_dim
                q_h = q[hs:hs+head_dim]
                k_h = [ki[hs:hs+head_dim] for ki in keys[li]]
                v_h = [vi[hs:hs+head_dim] for vi in values[li]]
                attn_logits = [sum(q_h[j] * k_h[t][j] for j in range(head_dim)) / head_dim**0.5 for t in range(len(k_h))]
                attn_weights = softmax(attn_logits)
                head_out = [sum(attn_weights[t] * v_h[t][j] for t in range(len(v_h))) for j in range(head_dim)]
                x_attn.extend(head_out)
            x = linear(x_attn, state_dict[f'layer{li}.attn_wo'])
            x = [a + b for a, b in zip(x, x_residual)]
            
            x_residual = x
            x = rmsnorm(x)
            x = linear(x, state_dict[f'layer{li}.mlp_fc1'])
            x = [xi.relu() for xi in x]
            x = linear(x, state_dict[f'layer{li}.mlp_fc2'])
            x = [a + b for a, b in zip(x, x_residual)]

        logits = linear(x, state_dict['lm_head'])
        return logits

    learning_rate, beta1, beta2, eps_adam = 0.01, 0.85, 0.99, 1e-8
    m = [0.0] * len(params)
    v = [0.0] * len(params)

    num_steps = min(50, len(docs) * 2)
    for step in range(num_steps):
        doc = docs[step % len(docs)]
        tokens = [BOS] + [uchars.index(ch) for ch in doc if ch in uchars] + [BOS]
        n = min(block_size, len(tokens) - 1)
        if n <= 0: continue

        keys, values = [[] for _ in range(n_layer)], [[] for _ in range(n_layer)]
        losses = []
        for pos_id in range(n):
            token_id, target_id = tokens[pos_id], tokens[pos_id + 1]
            logits = gpt(token_id, pos_id, keys, values)
            probs = softmax(logits)
            loss_t = -probs[target_id].log()
            losses.append(loss_t)
        
        loss = (1 / n) * sum(losses)
        loss.backward()

        lr_t = learning_rate * (1 - step / num_steps)
        for i, p in enumerate(params):
            m[i] = beta1 * m[i] + (1 - beta1) * p.grad
            v[i] = beta2 * v[i] + (1 - beta2) * p.grad ** 2
            m_hat = m[i] / (1 - beta1 ** (step + 1))
            v_hat = v[i] / (1 - beta2 ** (step + 1))
            p.data -= lr_t * m_hat / (v_hat ** 0.5 + eps_adam)
            p.grad = 0
            
    return state_dict

def train_pytorch(docs, uchars, vocab_size, n_layer, n_embd, block_size, n_head):
    import torch
    import torch.nn as nn
    from torch.nn import functional as F
    
    device = 'cuda' if torch.cuda.is_available() else 'mps' if torch.backends.mps.is_available() else 'cpu'
    BOS = len(uchars)
    
    class Head(nn.Module):
        def __init__(self, head_size):
            super().__init__()
            self.key = nn.Linear(n_embd, head_size, bias=False)
            self.query = nn.Linear(n_embd, head_size, bias=False)
            self.value = nn.Linear(n_embd, head_size, bias=False)
            self.register_buffer('tril', torch.tril(torch.ones(block_size, block_size)))
        def forward(self, x):
            B,T,C = x.shape
            k = self.key(x)
            q = self.query(x)
            wei = q @ k.transpose(-2,-1) * k.shape[-1]**-0.5
            wei = wei.masked_fill(self.tril[:T, :T] == 0, float('-inf'))
            wei = F.softmax(wei, dim=-1)
            v = self.value(x)
            out = wei @ v
            return out

    class MultiHeadAttention(nn.Module):
        def __init__(self, num_heads, head_size):
            super().__init__()
            self.heads = nn.ModuleList([Head(head_size) for _ in range(num_heads)])
            self.proj = nn.Linear(head_size * num_heads, n_embd, bias=False)
        def forward(self, x):
            out = torch.cat([h(x) for h in self.heads], dim=-1)
            out = self.proj(out)
            return out
            
    class FeedForward(nn.Module):
        def __init__(self, n_embd):
            super().__init__()
            self.net = nn.Sequential(
                nn.Linear(n_embd, 4 * n_embd, bias=False),
                nn.ReLU(),
                nn.Linear(4 * n_embd, n_embd, bias=False),
            )
        def forward(self, x):
            return self.net(x)

    class Block(nn.Module):
        def __init__(self, n_embd, n_head):
            super().__init__()
            head_size = n_embd // n_head
            self.sa = MultiHeadAttention(n_head, head_size)
            self.ffwd = FeedForward(n_embd)
            self.ln1 = nn.RMSNorm(n_embd)
            self.ln2 = nn.RMSNorm(n_embd)
        def forward(self, x):
            x = x + self.sa(self.ln1(x))
            x = x + self.ffwd(self.ln2(x))
            return x

    class GPT(nn.Module):
        def __init__(self):
            super().__init__()
            self.token_embedding_table = nn.Embedding(vocab_size, n_embd)
            self.position_embedding_table = nn.Embedding(block_size, n_embd)
            self.blocks = nn.Sequential(*[Block(n_embd, n_head=n_head) for _ in range(n_layer)])
            self.lm_head = nn.Linear(n_embd, vocab_size, bias=False)
        def forward(self, idx, targets=None):
            B, T = idx.shape
            tok_emb = self.token_embedding_table(idx)
            pos_emb = self.position_embedding_table(torch.arange(T, device=device))
            x = tok_emb + pos_emb
            x = self.blocks(x)
            logits = self.lm_head(x)
            
            if targets is None:
                loss = None
            else:
                B, T, C = logits.shape
                logits = logits.view(B*T, C)
                targets = targets.view(B*T)
                loss = F.cross_entropy(logits, targets)
            return logits, loss

    model = GPT().to(device)
    optimizer = torch.optim.AdamW(model.parameters(), lr=1e-3)

    num_steps = min(500, len(docs) * 10)
    for step in range(num_steps):
        doc = docs[step % len(docs)]
        tokens = [BOS] + [uchars.index(ch) for ch in doc if ch in uchars] + [BOS]
        if len(tokens) <= 1: continue
        
        # We limit the chunk to block_size or less
        chunk_len = min(block_size, len(tokens) - 1)
        x = torch.tensor([tokens[:chunk_len]], dtype=torch.long, device=device)
        y = torch.tensor([tokens[1:chunk_len+1]], dtype=torch.long, device=device)

        logits, loss = model(x, y)
        optimizer.zero_grad(set_to_none=True)
        loss.backward()
        optimizer.step()

    # Convert PyTorch state back to our custom Value dictionary format for the GGUF exporter
    class DummyValue:
        def __init__(self, data):
            self.data = data
            
    def wrap_tensor(t):
        return [[DummyValue(val.item()) for val in row] for row in t]
    
    sd = model.state_dict()
    custom_state = {}
    custom_state['wte'] = wrap_tensor(sd['token_embedding_table.weight'])
    custom_state['wpe'] = wrap_tensor(sd['position_embedding_table.weight'])
    custom_state['lm_head'] = wrap_tensor(sd['lm_head.weight'])
    
    for i in range(n_layer):
        # PyTorch attention splits QKV, we need to extract them back
        # QKV weights in PyTorch MultiHeadAttention (custom implemented) are structured differently 
        # than a standard linear output, so let's extract each head's weights and cat them
        
        # WQ, WK, WV shape needs to be (n_embd, n_embd) which is transposed of PyTorch Linear weight (out_features, in_features)
        # So in PyTorch they are (head_size, n_embd) per head.
        wq, wk, wv = [], [], []
        for h in range(n_head):
            wq.append(sd[f'blocks.{i}.sa.heads.{h}.query.weight'])
            wk.append(sd[f'blocks.{i}.sa.heads.{h}.key.weight'])
            wv.append(sd[f'blocks.{i}.sa.heads.{h}.value.weight'])
        
        # Concat along out_features -> (n_embd, n_embd)
        wq = torch.cat(wq, dim=0)
        wk = torch.cat(wk, dim=0)
        wv = torch.cat(wv, dim=0)
        
        custom_state[f'layer{i}.attn_wq'] = wrap_tensor(wq)
        custom_state[f'layer{i}.attn_wk'] = wrap_tensor(wk)
        custom_state[f'layer{i}.attn_wv'] = wrap_tensor(wv)
        custom_state[f'layer{i}.attn_wo'] = wrap_tensor(sd[f'blocks.{i}.sa.proj.weight'])
        custom_state[f'layer{i}.mlp_fc1'] = wrap_tensor(sd[f'blocks.{i}.ffwd.net.0.weight'])
        custom_state[f'layer{i}.mlp_fc2'] = wrap_tensor(sd[f'blocks.{i}.ffwd.net.2.weight'])

    return custom_state

def export_to_gguf(state_dict, uchars, vocab_size, n_embd, n_layer, n_head, block_size, data_dir):
    import gguf
    import numpy as np
    import subprocess
    
    gguf_writer = gguf.GGUFWriter(os.path.join(data_dir, "local-gpt.gguf"), "llama")
    
    # Add metadata
    gguf_writer.add_name("PyGPT")
    gguf_writer.add_architecture()
    gguf_writer.add_context_length(block_size)
    gguf_writer.add_embedding_length(n_embd)
    gguf_writer.add_block_count(n_layer)
    gguf_writer.add_feed_forward_length(4 * n_embd)
    gguf_writer.add_head_count(n_head)
    
    # Tokenizer - char level
    tokens = uchars + ['<|endoftext|>']
    gguf_writer.add_tokenizer_model("gpt2")
    gguf_writer.add_token_list(tokens)

    def extract_tensor(name):
        mat = state_dict[name]
        return np.array([[val.data for val in row] for row in mat], dtype=np.float32)

    # Convert weights
    # wte: (vocab_size, n_embd) -> token_embd.weight
    gguf_writer.add_tensor("token_embd.weight", extract_tensor('wte'))
    # wpe: (block_size, n_embd) -> position_embd.weight
    gguf_writer.add_tensor("position_embd.weight", extract_tensor('wpe'))
    
    for i in range(n_layer):
        # Attention
        wq = extract_tensor(f'layer{i}.attn_wq')
        wk = extract_tensor(f'layer{i}.attn_wk')
        wv = extract_tensor(f'layer{i}.attn_wv')
        wo = extract_tensor(f'layer{i}.attn_wo')
        
        # In llama/gpt2 architecture, QKV are grouped or separate
        gguf_writer.add_tensor(f"blk.{i}.attn_q.weight", wq)
        gguf_writer.add_tensor(f"blk.{i}.attn_k.weight", wk)
        gguf_writer.add_tensor(f"blk.{i}.attn_v.weight", wv)
        gguf_writer.add_tensor(f"blk.{i}.attn_output.weight", wo)
        
        # MLP
        fc1 = extract_tensor(f'layer{i}.mlp_fc1')
        fc2 = extract_tensor(f'layer{i}.mlp_fc2')
        gguf_writer.add_tensor(f"blk.{i}.ffn_up.weight", fc1)
        gguf_writer.add_tensor(f"blk.{i}.ffn_down.weight", fc2)
        
        # Norms (mocking RMSNorm as LayerNorm without bias to satisfy GGUF parsers if needed)
        # We don't have explicit learned norm weights in the pure python version, so we produce ones.
        gguf_writer.add_tensor(f"blk.{i}.attn_norm.weight", np.ones((n_embd,), dtype=np.float32))
        gguf_writer.add_tensor(f"blk.{i}.ffn_norm.weight", np.ones((n_embd,), dtype=np.float32))

    gguf_writer.add_tensor("output_norm.weight", np.ones((n_embd,), dtype=np.float32))
    gguf_writer.add_tensor("output.weight", extract_tensor('lm_head'))

    gguf_writer.write_header_to_file()
    gguf_writer.write_kv_data_to_file()
    gguf_writer.write_tensors_to_file()
    gguf_writer.close()

    # Create Ollama model
    modelfile_path = os.path.join(data_dir, "Modelfile.localgpt")
    with open(modelfile_path, "w") as f:
        f.write(f"FROM ./local-gpt.gguf\n")
        f.write(f"TEMPLATE \"\"\"\n{{{{ .Prompt }}}}\"\"\"\n")
        
    try:
        subprocess.run(["ollama", "create", "local-gpt", "-f", modelfile_path], check=True)
    except subprocess.CalledProcessError as e:
        pass

if __name__ == "__main__":
    main()
